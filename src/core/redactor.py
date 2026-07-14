# Copyright Amazon.com, Inc. or its affiliates. All Rights Reserved.
# SPDX-License-Identifier: Apache-2.0

"""
Core redaction logic for Step 3 (batch mode).
Routes to correct processor based on file type, handles mapping filtering and markers.
"""

import os
import io
import logging
import tempfile

logger = logging.getLogger(__name__)

# Extension → file type category
TEXT_EXTS = {".txt"}
JSON_EXTS = {".json"}
CSV_EXTS = {".csv"}
EXCEL_EXTS = {".xlsx"}
WORD_EXTS = {".docx"}
PDF_EXTS = {".pdf"}
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp"}


def build_file_mapping(detections, full_mapping):
    """Filter full synthetic mapping to only PII found in this file."""
    file_pii = {d.get("content", "").strip() for d in detections if d.get("content")}
    return {orig: syn for orig, syn in full_mapping.items() if orig in file_pii}


def build_blackout_mapping(detections):
    """Build mapping that replaces all PII with [REDACTED]."""
    return {
        d.get("content", "").strip(): "[REDACTED]"
        for d in detections
        if d.get("content", "").strip()
    }


def apply_text_markers(mapping):
    """Wrap synthetic values with *** markers."""
    return {orig: f"***{syn}***" for orig, syn in mapping.items()}


def redact_file(
    s3_client,
    source_bucket,
    source_key,
    output_bucket,
    folder_path,
    pii_mapping,
    detections,
    file_type,
    config,
    bedrock_runtime=None,
    job_id=None,
    output_prefix=None,
):
    """Route to correct redaction function based on file extension.

    Returns: (output_s3_key, unique_replaced_count, found_originals_set)
    """
    ext = os.path.splitext(source_key)[1].lower()
    markers = config.get("redaction", {}).get("markers", {})
    mode = config.get("redaction", {}).get("mode", "synthetic")

    if ext in TEXT_EXTS or ext in JSON_EXTS:
        if markers.get("text", False) and mode == "synthetic":
            pii_mapping = apply_text_markers(pii_mapping)
        return _redact_text_based(
            s3_client,
            source_bucket,
            source_key,
            output_bucket,
            folder_path,
            pii_mapping,
            detections=detections,
        )

    elif ext in CSV_EXTS:
        highlight = markers.get("tabular", False) and mode == "synthetic"
        return _redact_csv(
            s3_client,
            source_bucket,
            source_key,
            output_bucket,
            folder_path,
            pii_mapping,
            highlight,
        )

    elif ext in EXCEL_EXTS:
        highlight = markers.get("tabular", False) and mode == "synthetic"
        return _redact_excel(
            s3_client,
            source_bucket,
            source_key,
            output_bucket,
            folder_path,
            pii_mapping,
            highlight,
        )

    elif ext in WORD_EXTS:
        highlight = markers.get("word", False) and mode == "synthetic"
        return _redact_word(
            s3_client,
            source_bucket,
            source_key,
            output_bucket,
            folder_path,
            pii_mapping,
            highlight,
        )

    elif ext in PDF_EXTS:
        if file_type == "pdf_image":
            return _redact_pdf_image(
                s3_client,
                source_bucket,
                source_key,
                output_bucket,
                folder_path,
                pii_mapping,
                detections,
                config,
                bedrock_runtime,
                job_id=job_id,
                output_prefix=output_prefix,
            )
        else:
            return _redact_text_based(
                s3_client,
                source_bucket,
                source_key,
                output_bucket,
                folder_path,
                pii_mapping,
                is_pdf=True,
                detections=detections,
            )

    elif ext in IMAGE_EXTS:
        from helpers.model_config_helper import get_show_bounding_boxes

        show_boxes = get_show_bounding_boxes(config)
        return _redact_image(
            s3_client,
            source_bucket,
            source_key,
            output_bucket,
            folder_path,
            pii_mapping,
            detections,
            config,
            show_boxes,
        )

    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _redact_text_based(
    s3_client,
    source_bucket,
    source_key,
    output_bucket,
    folder_path,
    pii_mapping,
    is_pdf=False,
    detections=None,
):
    """Redact text/json/pdf-text files using text_replacer."""
    from core.text_replacer import replace_pii_in_text

    if is_pdf:
        from pypdf import PdfReader

        obj = s3_client.get_object(Bucket=source_bucket, Key=source_key)
        pdf_bytes = obj["Body"].read()
        reader = PdfReader(io.BytesIO(pdf_bytes))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    else:
        obj = s3_client.get_object(Bucket=source_bucket, Key=source_key)
        text = obj["Body"].read().decode("utf-8", errors="ignore")

    redacted, found, _, _, occurrence_counts = replace_pii_in_text(text, pii_mapping)

    # Set extra occurrences on detections for report
    # Only set on first detection of each value to avoid double-counting
    if detections:
        from collections import Counter

        det_counts = Counter(d.get("content", "") for d in detections)
        seen = set()
        for det in detections:
            content = det.get("content", "")
            if content in seen:
                continue
            seen.add(content)
            actual = occurrence_counts.get(content, 0)
            llm = det_counts.get(content, 0)
            if actual > llm:
                det["_extra_occurrences"] = actual - llm

    filename = os.path.basename(source_key)
    if is_pdf:
        filename = os.path.splitext(filename)[0] + ".txt"
    output_key = f"{folder_path}{filename}"

    content_type = "application/json" if source_key.endswith(".json") else "text/plain"
    s3_client.put_object(
        Body=redacted.encode("utf-8"),
        Bucket=output_bucket,
        Key=output_key,
        ContentType=content_type,
    )

    logger.info(f"[TEXT] replaced={len(found)}/{len(pii_mapping)}")
    return output_key, len(found), found


def _redact_csv(
    s3_client,
    source_bucket,
    source_key,
    output_bucket,
    folder_path,
    pii_mapping,
    highlight=False,
):
    """Redact CSV. If highlight, convert to xlsx with yellow cells."""
    from core.text_replacer import replace_pii_in_text

    obj = s3_client.get_object(Bucket=source_bucket, Key=source_key)
    text = obj["Body"].read().decode("utf-8", errors="ignore")

    if highlight:
        import csv
        import openpyxl
        from openpyxl.styles import PatternFill

        yellow = PatternFill(
            start_color="FFFF00", end_color="FFFF00", fill_type="solid"
        )
        wb = openpyxl.Workbook()
        ws = wb.active
        total_found = set()
        for r_idx, row in enumerate(csv.reader(io.StringIO(text)), 1):
            for c_idx, val in enumerate(row, 1):
                new_val, found, _, _, _ = replace_pii_in_text(val, pii_mapping)
                cell = ws.cell(row=r_idx, column=c_idx, value=new_val)
                if found:
                    cell.fill = yellow
                    total_found.update(found)

        filename = os.path.splitext(os.path.basename(source_key))[0]
        output_key = f"{folder_path}{filename}.xlsx"
        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        s3_client.put_object(Body=buf.getvalue(), Bucket=output_bucket, Key=output_key)
        logger.info(f"[CSV→XLSX] replaced={len(total_found)} highlight=True")
        return output_key, len(total_found), total_found
    else:
        redacted, found, _, _, _ = replace_pii_in_text(text, pii_mapping)
        filename = os.path.basename(source_key)
        output_key = f"{folder_path}{filename}"
        s3_client.put_object(
            Body=redacted.encode("utf-8"), Bucket=output_bucket, Key=output_key
        )
        logger.info(f"[CSV] replaced={len(found)}/{len(pii_mapping)}")
        return output_key, len(found), found


def _redact_excel(
    s3_client,
    source_bucket,
    source_key,
    output_bucket,
    folder_path,
    pii_mapping,
    highlight=False,
):
    """Redact xlsx cell by cell."""
    from core.text_replacer import replace_pii_in_text
    import openpyxl
    from openpyxl.styles import PatternFill

    yellow = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
    filename = os.path.basename(source_key)
    fd, temp_path = tempfile.mkstemp(suffix=f"_{filename}", dir=tempfile.gettempdir())
    os.close(fd)
    s3_client.download_file(source_bucket, source_key, temp_path)

    wb = openpyxl.load_workbook(temp_path)
    total_found = set()
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    new_val, found, _, _, _ = replace_pii_in_text(
                        str(cell.value), pii_mapping
                    )
                    if found:
                        cell.value = new_val
                        total_found.update(found)
                        if highlight:
                            cell.fill = yellow

    output_key = f"{folder_path}{filename}"
    fd2, out_path = tempfile.mkstemp(
        suffix=f"_redacted_{filename}", dir=tempfile.gettempdir()
    )
    os.close(fd2)
    wb.save(out_path)
    s3_client.upload_file(out_path, output_bucket, output_key)
    os.remove(temp_path)
    os.remove(out_path)

    logger.info(
        f"[XLSX] replaced={len(total_found)}/{len(pii_mapping)} highlight={highlight}"
    )
    return output_key, len(total_found), total_found


def _redact_word(
    s3_client,
    source_bucket,
    source_key,
    output_bucket,
    folder_path,
    pii_mapping,
    highlight=False,
):
    """Redact docx using existing replace_pii_in_word + optional highlight."""
    from processors.word_processor import replace_pii_in_word

    filename = os.path.basename(source_key)
    fd, temp_path = tempfile.mkstemp(suffix=f"_{filename}", dir=tempfile.gettempdir())
    os.close(fd)
    fd2, out_path = tempfile.mkstemp(
        suffix=f"_redacted_{filename}", dir=tempfile.gettempdir()
    )
    os.close(fd2)
    s3_client.download_file(source_bucket, source_key, temp_path)

    result = replace_pii_in_word(temp_path, out_path, pii_mapping)

    if highlight and result.get("found_originals"):
        _highlight_word_runs(out_path, result["found_originals"], pii_mapping)

    output_key = f"{folder_path}{filename}"
    s3_client.upload_file(out_path, output_bucket, output_key)
    os.remove(temp_path)
    os.remove(out_path)

    replaced = len(result.get("found_originals", set()))
    logger.info(f"[DOCX] replaced={replaced}/{len(pii_mapping)} highlight={highlight}")
    return output_key, replaced, result.get("found_originals", set())


def _highlight_word_runs(docx_path, found_originals, pii_mapping):
    """Add yellow highlight only to the synthetic replacement text, not the full run."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from lxml import etree
    import copy

    doc = Document(docx_path)
    syn_values = {pii_mapping[orig] for orig in found_originals if orig in pii_mapping}

    def _split_and_highlight(run, parent_element):
        """Split a run so only synthetic values get highlighted. Returns True if split."""
        text = run.text or ""
        if not text:
            return False

        # Find all synthetic value positions in this run
        segments = []
        remaining = text
        while remaining:
            earliest_match = None
            earliest_idx = len(remaining)
            for sv in syn_values:
                idx = remaining.find(sv)
                if idx != -1 and idx < earliest_idx:
                    earliest_idx = idx
                    earliest_match = sv
            if earliest_match is None:
                segments.append((remaining, False))
                break
            if earliest_idx > 0:
                segments.append((remaining[:earliest_idx], False))
            segments.append((earliest_match, True))
            remaining = remaining[earliest_idx + len(earliest_match) :]

        if not any(hl for _, hl in segments):
            return False
        if len(segments) == 1 and segments[0][1]:
            # Entire run is the synthetic value — just highlight it
            rpr = run._element.get_or_add_rPr()
            hl = etree.SubElement(rpr, qn("w:highlight"))
            hl.set(qn("w:val"), "yellow")
            return True

        # Split into multiple runs
        run_parent = run._element.getparent()
        run_idx = list(run_parent).index(run._element)
        for i, (seg_text, should_hl) in enumerate(segments):
            if i == 0:
                run.text = seg_text
                if should_hl:
                    rpr = run._element.get_or_add_rPr()
                    hl = etree.SubElement(rpr, qn("w:highlight"))
                    hl.set(qn("w:val"), "yellow")
            else:
                new_run = copy.deepcopy(run._element)
                # Set text on the w:t element
                t_elem = new_run.find(qn("w:t"))
                if t_elem is None:
                    t_elem = OxmlElement("w:t")
                    new_run.append(t_elem)
                t_elem.text = seg_text
                t_elem.set(qn("xml:space"), "preserve")
                # Remove or add highlight
                rpr = new_run.find(qn("w:rPr"))
                if rpr is not None:
                    for old_hl in rpr.findall(qn("w:highlight")):
                        rpr.remove(old_hl)
                if should_hl:
                    if rpr is None:
                        rpr = OxmlElement("w:rPr")
                        new_run.insert(0, rpr)
                    hl = etree.SubElement(rpr, qn("w:highlight"))
                    hl.set(qn("w:val"), "yellow")
                run_parent.insert(run_idx + i, new_run)
        return True

    for para in doc.paragraphs:
        for run in list(para.runs):
            if any(sv in (run.text or "") for sv in syn_values):
                _split_and_highlight(run, para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in list(para.runs):
                        if any(sv in (run.text or "") for sv in syn_values):
                            _split_and_highlight(run, para)

    doc.save(docx_path)


def _redact_pdf_image(
    s3_client,
    source_bucket,
    source_key,
    output_bucket,
    folder_path,
    pii_mapping,
    detections,
    config,
    bedrock_runtime,
    job_id=None,
    output_prefix=None,
):
    """Redact image-based PDF using bounding boxes via existing redact_pdf."""
    from redaction.pdf_redactor import redact_pdf

    filename = os.path.splitext(os.path.basename(source_key))[0]
    obj = s3_client.get_object(Bucket=source_bucket, Key=source_key)
    pdf_bytes = obj["Body"].read()

    fd, temp_in = tempfile.mkstemp(
        suffix=f"_in_{filename}.pdf", dir=tempfile.gettempdir()
    )
    os.close(fd)
    fd2, temp_out = tempfile.mkstemp(
        suffix=f"_out_{filename}.pdf", dir=tempfile.gettempdir()
    )
    os.close(fd2)
    fd3, temp_summary = tempfile.mkstemp(
        suffix=f"_sum_{filename}.json", dir=tempfile.gettempdir()
    )
    os.close(fd3)

    with open(temp_in, "wb") as f:
        f.write(pdf_bytes)

    try:
        from helpers.model_config_helper import get_show_bounding_boxes

        redaction_config = {
            **config.get("redaction", {}),
            "show_bounding_boxes": get_show_bounding_boxes(config),
            "bounding_box_color": config.get("validation", {}).get(
                "bounding_box_color", [1, 0, 0]
            ),
            "dpi": config.get("performance", {}).get("dpi", 300),
        }
        _rpfx = f"{output_prefix}/" if output_prefix else ""
        _safe = os.path.basename(source_key).replace(".", "_")
        if output_prefix:
            _render_pfx = f"{_rpfx}intermediate/renders/{_safe}/"
        else:
            _render_pfx = f"intermediate/{_safe}/renders/"
        success = redact_pdf(
            temp_in,
            temp_out,
            detections,
            pii_mapping,
            temp_summary,
            None,
            None,
            bedrock_runtime=bedrock_runtime,
            redaction_config=redaction_config,
            s3_bucket=output_bucket,
            s3_render_prefix=_render_pfx,
        )
        if not success:
            raise RuntimeError("PDF image redaction failed")

        output_key = f"{folder_path}{filename}.pdf"
        with open(temp_out, "rb") as f:
            s3_client.put_object(Body=f.read(), Bucket=output_bucket, Key=output_key)

        redacted_count = len(
            [
                d
                for d in detections
                if d.get("bounding_box") and not d.get("_font_skipped")
            ]
        )
        logger.info(f"[PDF-IMAGE] redacted={redacted_count}")
        return (
            output_key,
            redacted_count,
            {
                d.get("content", "")
                for d in detections
                if d.get("bounding_box") and not d.get("_font_skipped")
            },
        )
    finally:
        for p in [temp_in, temp_out, temp_summary]:
            if os.path.exists(p):
                os.remove(p)


def _redact_image(
    s3_client,
    source_bucket,
    source_key,
    output_bucket,
    folder_path,
    pii_mapping,
    detections,
    config,
    show_boxes=False,
):
    """Redact image file using bounding boxes."""
    from PIL import ImageDraw
    from processors.image_processor import (
        load_image_from_s3,
        draw_synthetic_text_on_image,
        MULTI_PAGE_FORMATS,
    )

    images = load_image_from_s3(s3_client, source_bucket, source_key, config)
    ext = os.path.splitext(source_key)[1].lower()
    mode = config.get("redaction", {}).get("mode", "synthetic")
    raw_color = config.get("validation", {}).get("bounding_box_color", [255, 0, 0])
    box_color = tuple(int(c * 255) if c <= 1.0 else int(c) for c in raw_color)

    # Group detections by page
    page_dets = {}
    for d in detections:
        page_dets.setdefault(d.get("page_num", 1), []).append(d)

    redacted_images = []
    total_redacted = 0

    for page_num, img in enumerate(images, start=1):
        dets = page_dets.get(page_num, [])
        if not dets:
            redacted_images.append(img)
            continue

        img_r = img.copy()
        draw = ImageDraw.Draw(img_r)

        for det in dets:
            bbox = det.get("bounding_box")
            if not bbox:
                continue

            original = det.get("content", "")
            synthetic = pii_mapping.get(original, original)

            left = max(0, int(bbox["left"]))
            top = max(0, int(bbox["top"]))
            right = min(img.width, int(bbox["left"] + bbox["width"]))
            bottom = min(img.height, int(bbox["top"] + bbox["height"]))

            if mode == "blackout":
                draw.rectangle([left, top, right, bottom], fill=(0, 0, 0))
            else:
                draw.rectangle([left, top, right, bottom], fill=(255, 255, 255))
                draw_synthetic_text_on_image(
                    draw, synthetic, left, top, right - left, bottom - top
                )

            if show_boxes:
                draw.rectangle([left, top, right, bottom], outline=box_color, width=3)

            total_redacted += 1

        redacted_images.append(img_r)

    # Save
    filename = os.path.splitext(os.path.basename(source_key))[0]
    output_key = f"{folder_path}{filename}{ext}"

    buf = io.BytesIO()
    if ext in MULTI_PAGE_FORMATS and len(redacted_images) > 1:
        redacted_images[0].save(
            buf,
            format="TIFF",
            save_all=True,
            append_images=redacted_images[1:],
            compression="tiff_deflate",
        )
    elif ext in [".tiff", ".tif"]:
        redacted_images[0].save(buf, format="TIFF", compression="tiff_deflate")
    else:
        fmt = "PNG" if ext == ".png" else "JPEG"
        redacted_images[0].save(buf, format=fmt)
    buf.seek(0)

    s3_client.put_object(Body=buf.getvalue(), Bucket=output_bucket, Key=output_key)
    logger.info(f"[IMAGE] redacted={total_redacted} pages={len(images)}")
    return (
        output_key,
        total_redacted,
        {d.get("content", "") for d in detections if d.get("bounding_box")},
    )
