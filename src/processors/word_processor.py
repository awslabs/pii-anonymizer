# Copyright Amazon.com, Inc. or its affiliates. All Rights Reserved.
# SPDX-License-Identifier: Apache-2.0

"""
Word-based PII Processor Module

3-step pipeline: detect PII per page concurrently, batch generate synthetic
replacements, replace in document preserving formatting via run manipulation.
"""

import os
import tempfile
import logging

from docx import Document

from helpers.model_config_helper import (
    get_inference_config_from_yaml,
    get_concurrency_config,
)
from core.synthetic_pii_generator import batch_generate_synthetic_pii
from helpers.threaded_detector import run_threaded_pii_detection
from helpers.token_tracker import TokenTracker

logger = logging.getLogger(__name__)

WPML_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def detect_pii_word(source_bucket, source_key, config, bedrock_runtime, s3_client):
    """Detect PII in a Word file (Step 1 only)."""
    filename = os.path.splitext(os.path.basename(source_key))[0]
    model_id = config["model"]["id"]
    model_provider = config["model"]["provider"]
    inference_config = get_inference_config_from_yaml(config)
    cc = get_concurrency_config(config)
    tracker = TokenTracker(model_id)

    temp_fd, temp_path = tempfile.mkstemp(
        suffix=f"_{filename}.docx", dir=tempfile.gettempdir()
    )
    os.close(temp_fd)
    s3_client.download_file(source_bucket, source_key, temp_path)
    logger.info(f"Downloaded Word file: {source_key}")

    from validation.document_validator import validate_word, DocumentValidationError

    try:
        validate_word(temp_path, config)
    except DocumentValidationError as e:
        os.remove(temp_path)
        raise ValueError(f"Word validation failed: {e}")

    doc = Document(temp_path)
    chunks = _build_page_chunks(doc)
    os.remove(temp_path)

    all_detections, failed_chunks, _ = run_threaded_pii_detection(
        chunks,
        model_id,
        model_provider,
        bedrock_runtime,
        inference_config,
        cc["max_workers"],
        label="Page",
        token_tracker=tracker,
        config=config,
    )

    return {
        "source_key": source_key,
        "file_type": "docx",
        "detections": [
            {
                "content": d["content"],
                "type": d["type"],
                "confidence": d.get("confidence", 0.0),
            }
            for d in all_detections
        ],
        "failed_chunks": failed_chunks,
        "token_usage": tracker.summary(),
    }


def _build_page_map(doc):
    """Build a map of paragraph text → page number using lastRenderedPageBreak tags.
    Returns None if the doc has no page break markers."""
    body = doc.element.body
    has_page_breaks = any(
        elem.tag == f"{WPML_NS}lastRenderedPageBreak"
        or (elem.tag == f"{WPML_NS}br" and elem.get(f"{WPML_NS}type") == "page")
        for elem in body.iter()
    )
    if not has_page_breaks:
        logger.info("No page break markers found — page numbers unavailable")
        return None

    page = 1
    page_map = {}
    for para in body.iter(f"{WPML_NS}p"):
        for elem in para.iter():
            if elem.tag == f"{WPML_NS}lastRenderedPageBreak" or (
                elem.tag == f"{WPML_NS}br" and elem.get(f"{WPML_NS}type") == "page"
            ):
                page += 1
        text = "".join(node.text or "" for node in para.iter(f"{WPML_NS}t")).strip()
        if text:
            page_map[text] = page
    return page_map


def _build_page_chunks(doc):
    """Build {page_id: text} chunks for concurrent detection.
    Falls back to single chunk if no page breaks found."""
    body = doc.element.body
    page = 1
    pages = {}
    current_texts = []

    for para in body.iter(f"{WPML_NS}p"):
        for elem in para.iter():
            if elem.tag == f"{WPML_NS}lastRenderedPageBreak" or (
                elem.tag == f"{WPML_NS}br" and elem.get(f"{WPML_NS}type") == "page"
            ):
                if current_texts:
                    pages[f"Page {page}"] = "\n".join(current_texts)
                    current_texts = []
                page += 1
        text = "".join(node.text or "" for node in para.iter(f"{WPML_NS}t")).strip()
        if text:
            current_texts.append(text)

    if current_texts:
        pages[f"Page {page}"] = "\n".join(current_texts)

    # Include table text appended to last page
    table_texts = []
    for t in doc.tables:
        try:
            for row in t.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    table_texts.append(row_text)
        except Exception as e:
            logger.warning(f"Skipping malformed table: {e}")

    if table_texts:
        last_key = list(pages.keys())[-1] if pages else "Page 1"
        pages[last_key] = pages.get(last_key, "") + "\n" + "\n".join(table_texts)

    if not pages:
        all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if all_text.strip():
            pages["Page 1"] = all_text

    image_count = len(doc.inline_shapes)
    if image_count:
        logger.warning(
            f"Document contains {image_count} embedded image(s) — PII in images will NOT be detected"
        )

    return pages


def _replace_in_runs(runs, orig, syn):
    """Replace text across runs by combining, replacing, then redistributing."""
    combined = "".join(r.text for r in runs)
    if orig not in combined:
        return 0
    runs[0].text = combined.replace(orig, syn)
    for r in runs[1:]:
        r.text = ""
    return 1


def replace_pii_in_word(docx_path, output_path, pii_mapping):
    """Replace PII in Word document. Longest-first with placeholders to prevent corruption."""
    doc = Document(docx_path)
    found_originals = set()
    replacements = 0
    placeholders = {}
    ph_count = 0

    sorted_mapping = sorted(pii_mapping.items(), key=lambda x: len(x[0]), reverse=True)

    def _make_ph():
        nonlocal ph_count
        n = ph_count
        ph_count += 1
        return "\ue000" + "\ue001" * (n + 1) + "\ue000"

    # Pass 1: Replace originals with placeholders
    for orig, syn in sorted_mapping:
        ph = _make_ph()
        matched = False
        for para in doc.paragraphs:
            if orig in para.text and para.runs:
                count = _replace_in_runs(para.runs, orig, ph)
                replacements += count
                if count:
                    matched = True

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if orig in cell.text:
                        for para in cell.paragraphs:
                            if para.runs:
                                count = _replace_in_runs(para.runs, orig, ph)
                                replacements += count
                                if count:
                                    matched = True

        if matched:
            found_originals.add(orig)
            placeholders[ph] = syn

    # Pass 2: Swap placeholders with synthetic values
    for ph, syn in placeholders.items():
        for para in doc.paragraphs:
            if ph in para.text and para.runs:
                _replace_in_runs(para.runs, ph, syn)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if ph in cell.text:
                        for para in cell.paragraphs:
                            if para.runs:
                                _replace_in_runs(para.runs, ph, syn)

    not_found = len(pii_mapping) - len(found_originals)
    if not_found:
        missing = [k for k in pii_mapping if k not in found_originals]
        covered = [m for m in missing if any(m in r for r in found_originals)]
        # Build source text to check for hallucinated detections
        source_text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    source_text += "\n" + cell.text
        hallucinated = [m for m in missing if m not in covered and m not in source_text]
        truly_missing = len(missing) - len(covered) - len(hallucinated)
        if covered:
            logger.info(
                f"{len(covered)} PII skipped (already covered by longer replacement)"
            )
        if hallucinated:
            logger.info(
                f"{len(hallucinated)} PII skipped (LLM hallucinated, not in source text)"
            )
        if truly_missing:
            logger.warning(f"{truly_missing} PII not found in document")

    doc.save(output_path)
    logger.info(
        f"Replacement: {len(found_originals)}/{len(pii_mapping)} unique PII replaced, {not_found} not found"
    )
    return {
        "success": True,
        "replacements": replacements,
        "found_originals": found_originals,
    }


def _find_page_for_pii(pii_original, page_map):
    """Find which page a PII value appears on."""
    for text, page in page_map.items():
        if pii_original in text:
            return page
    return 0


def store_word_pii_mapping(
    detections,
    pii_mapping,
    filename,
    dynamodb_manager,
    page_map=None,
    found_originals=None,
    tracker=None,
):
    """Store Word PII mappings in DynamoDB."""
    try:
        detailed_pii_data = []
        found_set = found_originals or set()
        for det in detections:
            original = det["content"]
            synthetic = pii_mapping.get(original, "")
            if original in found_set:
                status = "text_replaced"
            elif any(original in f for f in found_set):
                status = "text_replaced"
                if not synthetic:
                    for f in found_set:
                        if original in f and f in pii_mapping:
                            synthetic = f"(covered by: {pii_mapping[f]})"
                            break
            else:
                status = "not_redacted"
            record = {
                "original": original,
                "synthetic": synthetic,
                "type": det.get("type", "UNKNOWN"),
                "confidence": det.get("confidence", 0),
                "source": "word",
                "replacement_status": status,
            }
            if status == "not_redacted":
                record["not_redacted_reason"] = "text_not_found_in_document"
            if page_map:
                page = _find_page_for_pii(original, page_map)
                if page:
                    record["page_num"] = page
            detailed_pii_data.append(record)

        token_usage = tracker.summary() if tracker else {}
        dynamodb_manager.store_pii_mapping(
            detailed_pii_data, filename, status="SUCCESS", token_usage=token_usage
        )
        if tracker:
            tracker.log_summary()
        replaced = sum(
            1 for d in detailed_pii_data if d["replacement_status"] == "text_replaced"
        )
        logger.info(
            f"Stored {len(detailed_pii_data)} Word PII mappings in DynamoDB ({replaced} replaced, {len(detailed_pii_data) - replaced} not found)"
        )
    except Exception as e:
        logger.warning(f"Failed to store Word PII mapping in DynamoDB: {e}")


def process_word_file(
    source_bucket,
    source_key,
    output_bucket,
    filename_without_ext,
    config,
    bedrock_runtime,
    dynamodb_manager,
    s3_client,
    folder_path="",
):
    """End-to-end Word file processing: detect per page, batch synthetic, replace."""
    try:
        model_id = config["model"]["id"]
        model_provider = config["model"]["provider"]
        inference_config = get_inference_config_from_yaml(config)
        cc = get_concurrency_config(config)
        tracker = TokenTracker(model_id)

        # Download from S3
        temp_fd, temp_path = tempfile.mkstemp(
            suffix=f"_{filename_without_ext}.docx", dir=tempfile.gettempdir()
        )
        os.close(temp_fd)
        s3_client.download_file(source_bucket, source_key, temp_path)
        logger.info(f"Downloaded Word file: {source_key}")

        # Validate Word document
        from validation.document_validator import validate_word, DocumentValidationError

        try:
            validate_word(temp_path, config)
            logger.info(f"Word validation passed for: {filename_without_ext}")
        except DocumentValidationError as e:
            logger.error(f"Word validation failed: {str(e)}")
            return {"success": False, "error": str(e), "pii_count": 0}

        # Step 1: Detect PII per page concurrently
        doc = Document(temp_path)
        chunks = _build_page_chunks(doc)
        page_map = _build_page_map(doc)
        logger.info(f"Word document: {len(chunks)} page(s)")

        all_detections, failed_chunks, raw_detections = run_threaded_pii_detection(
            chunks,
            model_id,
            model_provider,
            bedrock_runtime,
            inference_config,
            cc["max_workers"],
            label="Page",
            token_tracker=tracker,
            config=config,
        )
        logger.info(
            f"Total detections: {len(raw_detections)} raw, {len(all_detections)} after dedup"
        )

        if failed_chunks:
            logger.warning(f"Failed pages: {[f['chunk_id'] for f in failed_chunks]}")
            if dynamodb_manager:
                for fc in failed_chunks:
                    dynamodb_manager.store_pii_mapping(
                        [],
                        filename_without_ext,
                        status="FAILED_DETECTION",
                        error_message=f"Page '{fc['chunk_id']}': {fc['error']}",
                    )

        # Step 2: Batch generate synthetic replacements
        pii_mapping = {}
        if all_detections:
            pii_mapping = batch_generate_synthetic_pii(
                all_detections,
                model_id,
                model_provider,
                bedrock_runtime,
                config=config,
                token_tracker=tracker,
            )
            logger.info(f"Generated {len(pii_mapping)} synthetic replacements")

        # Step 3: Replace in document
        fd, output_path = tempfile.mkstemp(
            suffix=f"_redacted_{filename_without_ext}.docx", dir=tempfile.gettempdir()
        )
        os.close(fd)
        result = replace_pii_in_word(temp_path, output_path, pii_mapping)

        # Upload to S3
        s3_output_key = f"{folder_path}redacted_{filename_without_ext}.docx"
        s3_client.upload_file(output_path, output_bucket, s3_output_key)
        logger.info(f"Uploaded redacted Word file to: {s3_output_key}")

        # Create and upload summary JSON
        import json

        summary = {
            "input_file": f"s3://{source_bucket}/{source_key}",
            "output_file": f"s3://{output_bucket}/{s3_output_key}",
            "pii_count": len(all_detections),
            "pii_replaced": len(result["found_originals"]),
            "pii_not_found": len(pii_mapping) - len(result["found_originals"]),
            "pages_processed": len(chunks),
            "token_usage": tracker.summary() if tracker else {},
        }
        summary_key = f"{folder_path}redaction_summary_{filename_without_ext}.json"
        s3_client.put_object(
            Body=json.dumps(summary, indent=2),
            Bucket=output_bucket,
            Key=summary_key,
            ContentType="application/json",
        )
        logger.info(f"Uploaded summary to: {summary_key}")

        # Store in DynamoDB
        if dynamodb_manager and all_detections:
            store_word_pii_mapping(
                raw_detections,
                pii_mapping,
                filename_without_ext,
                dynamodb_manager,
                page_map,
                result["found_originals"],
                tracker,
            )

        # Cleanup
        os.remove(temp_path)
        os.remove(output_path)

        return {
            "success": True,
            "s3_output_file": s3_output_key,
            "pii_count": len(all_detections),
            "replacements": result["replacements"],
        }
    except Exception as e:
        logger.error(f"Error processing Word file: {e}", exc_info=True)
        if dynamodb_manager:
            dynamodb_manager.store_pii_mapping(
                [], filename_without_ext, status="FAILED", error_message=str(e)
            )
        return {"success": False, "error": str(e)}
